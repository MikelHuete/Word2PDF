import argparse
import os
import zipfile
import json
from datetime import datetime
from docx import Document

HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Informe de Inspección DOCX</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&display=swap" rel="stylesheet">
    <style>
        :root {{
            --primary: #2563eb;
            --bg: #f8fafc;
            --card: #ffffff;
            --text: #1e293b;
            --text-muted: #64748b;
            --border: #e2e8f0;
            --accent: #f1f5f9;
        }}
        body {{
            font-family: 'Inter', sans-serif;
            background-color: var(--bg);
            color: var(--text);
            line-height: 1.6;
            margin: 0;
            padding: 40px 20px;
        }}
        .container {{
            max-width: 1000px;
            margin: 0 auto;
        }}
        header {{
            margin-bottom: 40px;
            border-bottom: 2px solid var(--border);
            padding-bottom: 20px;
        }}
        h1 {{ font-weight: 700; color: var(--primary); margin: 0; }}
        .timestamp {{ color: var(--text-muted); font-size: 0.9rem; }}
        
        .section {{
            background: var(--card);
            border-radius: 12px;
            padding: 24px;
            margin-bottom: 24px;
            box-shadow: 0 4px 6px -1px rgb(0 0 0 / 0.1);
            border: 1px solid var(--border);
        }}
        h2 {{ border-left: 4px solid var(--primary); padding-left: 12px; font-size: 1.25rem; margin-top: 0; }}
        
        .grid {{ display: grid; grid-template-columns: repeat(auto-fill, minmax(300px, 1fr)); gap: 16px; }}
        .meta-item {{ border-bottom: 1px solid var(--border); padding: 8px 0; }}
        .meta-label {{ font-weight: 600; color: var(--text-muted); font-size: 0.85rem; text-transform: uppercase; }}
        
        .para {{ margin-bottom: 16px; padding: 12px; border-radius: 6px; background: var(--accent); }}
        .para-style {{ font-size: 0.75rem; color: var(--primary); font-weight: 600; margin-bottom: 4px; }}
        .run {{ display: inline-block; margin-right: 4px; }}
        
        table {{ width: 100%; border-collapse: collapse; margin-top: 12px; }}
        th, td {{ border: 1px solid var(--border); padding: 10px; text-align: left; }}
        th {{ background: var(--accent); font-weight: 600; }}
        
        .file-list {{ font-family: monospace; font-size: 0.85rem; background: #1e293b; color: #e2e8f0; padding: 16px; border-radius: 8px; overflow-x: auto; }}
        .file-item {{ margin-bottom: 4px; }}

        .tag {{ display: inline-block; padding: 2px 6px; border-radius: 4px; font-size: 0.7rem; font-weight: 700; margin-left: 4px; }}
        .tag-bold {{ background: #fee2e2; color: #991b1b; }}
        .tag-italic {{ background: #fef9c3; color: #854d0e; }}
        .tag-size {{ background: #dcfce7; color: #166534; }}
    </style>
</head>
<body>
    <div class="container">
        <header>
            <h1>Informe de Inspección DOCX</h1>
            <div class="timestamp">Generado el: {date} | Archivo: {filename}</div>
        </header>

        <div class="section">
            <h2>Metadatos</h2>
            <div class="grid">{metadata}</div>
        </div>

        <div class="section" id="content">
            <h2>Contenido y Estilos (Párrafos)</h2>
            {paragraphs}
        </div>

        <div class="section">
            <h2>Tablas Detectadas</h2>
            {tables}
        </div>

        <div class="section">
            <h2>Estructura Interna (ZIP Limit Testing)</h2>
            <p>Total de archivos internos: <strong>{file_count}</strong></p>
            <div class="file-list">{files}</div>
        </div>
        
        <div class="section">
            <h2>Medios y Archivos Incrustados</h2>
            <p>{media_msg}</p>
            <ul>{media_list}</ul>
        </div>
    </div>
</body>
</html>
"""

def get_report_data(docx_path, output_dir):
    data = {
        "filename": os.path.basename(docx_path),
        "date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "metadata": [],
        "paragraphs": [],
        "tables": [],
        "internal_files": [],
        "media_extracted": []
    }

    doc = Document(docx_path)
    
    # Metadata
    props = doc.core_properties
    for attr in ['author', 'category', 'comments', 'created', 'last_modified_by', 'modified', 'revision', 'title']:
        try:
            val = getattr(props, attr)
            if val:
                data["metadata"].append({"label": attr, "value": str(val)})
        except: pass

    # Paragraphs & Runs
    for para in doc.paragraphs:
        if para.text.strip():
            p_data = {"style": para.style.name, "runs": []}
            for run in para.runs:
                if run.text.strip():
                    r_data = {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "size": f"{run.font.size.pt}pt" if run.font.size else None
                    }
                    p_data["runs"].append(r_data)
            data["paragraphs"].append(p_data)

    # Tables
    for table in doc.tables:
        t_data = []
        for row in table.rows:
            t_data.append([cell.text.strip() for cell in row.cells])
        data["tables"].append(t_data)

    # ZIP Structure & Media
    media_dir = os.path.join(output_dir, "extracted_media")
    os.makedirs(media_dir, exist_ok=True)
    
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            for filename in z.namelist():
                info = z.getinfo(filename)
                data["internal_files"].append(f"{filename} ({info.file_size} bytes)")
                
                if filename.startswith('word/media/') or filename.startswith('word/embeddings/'):
                    z.extract(filename, media_dir)
                    data["media_extracted"].append(filename)
    except zipfile.BadZipFile:
        print(f"Error: {docx_path} no es un archivo DOCX (ZIP) válido.")
        exit(1)

    return data

def generate_html(data, report_path):
    meta_html = "".join([f'<div class="meta-item"><div class="meta-label">{m["label"]}</div><div>{m["value"]}</div></div>' for m in data["metadata"]])
    
    para_html = ""
    for p in data["paragraphs"]:
        para_html += f'<div class="para"><div class="para-style">{p["style"]}</div>'
        for r in p["runs"]:
            tags = ""
            if r["bold"]: tags += '<span class="tag tag-bold">B</span>'
            if r["italic"]: tags += '<span class="tag tag-italic">I</span>'
            if r["size"]: tags += f'<span class="tag tag-size">{r["size"]}</span>'
            para_html += f'<span class="run">{r["text"]}{tags}</span>'
        para_html += '</div>'

    table_html = ""
    for t_idx, t in enumerate(data["tables"]):
        table_html += f"<h3>Tabla {t_idx}</h3><table>"
        for r_idx, row in enumerate(t):
            table_html += "<tr>"
            for cell in row:
                tag = "th" if r_idx == 0 else "td"
                table_html += f"<{tag}>{cell}</{tag}>"
            table_html += "</tr>"
        table_html += "</table>"

    files_html = "".join([f'<div class="file-item">{f}</div>' for f in data["internal_files"]])
    
    media_msg = f"Se han extraído {len(data['media_extracted'])} archivos a la carpeta 'extracted_media'." if data['media_extracted'] else "No se encontraron medios incrustados."
    media_list = "".join([f"<li>{m}</li>" for m in data["media_extracted"]])

    report = HTML_TEMPLATE.format(
        date=data["date"],
        filename=data["filename"],
        metadata=meta_html,
        paragraphs=para_html,
        tables=table_html,
        file_count=len(data["internal_files"]),
        files=files_html,
        media_msg=media_msg,
        media_list=media_list
    )
    
    os.makedirs(os.path.dirname(os.path.abspath(report_path)), exist_ok=True)
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(report)
    
    print(f"\n[!] Informe generado exitosamente en: {report_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Inspecciona un archivo DOCX y genera un informe HTML detallado.")
    parser.add_argument("input", help="Ruta al archivo .docx a inspeccionar")
    parser.add_argument("-o", "--output", help="Ruta de salida para el informe HTML (por defecto report.html en el directorio del input)")
    
    args = parser.parse_args()
    
    target = args.input
    if not os.path.exists(target):
        print(f"Error: No se encontró el archivo {target}")
        exit(1)
        
    print(f"Procesando: {target}...")
    
    if args.output:
        # If output is a .html file, use its directory for results
        if args.output.lower().endswith(".html"):
            output_dir = os.path.dirname(os.path.abspath(args.output))
            report_out = args.output
        else:
            # Otherwise treat it as a directory
            output_dir = os.path.abspath(args.output)
            report_out = os.path.join(output_dir, "report.html")
    else:
        # Default output directory next to the input file
        output_dir = os.path.join(os.path.dirname(os.path.abspath(target)), "inspection_results")
        report_out = os.path.join(output_dir, "report.html")

    os.makedirs(output_dir, exist_ok=True)
    
    data = get_report_data(target, output_dir)
    generate_html(data, report_out)
